import os
import re
import base64
import shutil
from robot.api import ExecutionResult, ResultVisitor
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import io
from PIL import Image
from libraries.common.utility_helpers import PROJECT_ROOT


class ScreenshotCollector(ResultVisitor):
    def __init__(self):
        self.test_results = []
        self.current_test = None
        self.base64_screenshots = []
        self.current_keyword = None
        self.last_action_step = None

    def start_test(self, test):
        self.current_test = {
            'id': test.id,
            'name': test.name,
            'doc': test.doc,
            'status': test.status,
            'suites': self._get_parent_suites(test),
            'steps': [],
            'base64_screenshots': [],
            'start_time': test.starttime,
            'end_time': test.endtime,
            'duration': test.elapsedtime
        }
        self.base64_screenshots = []

    def end_test(self, test):
        self.current_test['base64_screenshots'] = self.base64_screenshots
        self.test_results.append(self.current_test)
        self.current_test = None
        self.current_keyword = None

    def start_keyword(self, keyword):
        if self.current_test is None:
            return

        step_info = {
            'name': keyword.name,
            'status': keyword.status,
            'messages': []
        }
        self.current_keyword = step_info

        for msg in keyword.messages:
            if "Screenshot captured successfully at:" in msg.message:
                description_match = re.search(r'description:\s*(.+?)(?:$|\s)', msg.message)
                if description_match:
                    description = description_match.group(1)
                    if description:
                        self.last_action_step = {
                            'name': self.current_keyword['name'],
                            'description': description
                        }

            if msg.html and "img src=" in msg.message:
                base64_match = re.search(r'img src="data:image/[^;]+;base64,([^"]+)"', msg.message)
                if base64_match:
                    screenshot_info = {
                        'data': base64_match.group(1),
                        'step': self.last_action_step if self.last_action_step else {'name': self.current_keyword['name']}
                    }
                    self.base64_screenshots.append(screenshot_info)

        if step_info['messages']:
            self.current_test['steps'].append(step_info)

    def _get_parent_suites(self, test):
        suites = []
        current = test.parent
        while current is not None and hasattr(current, 'name'):
            suites.append({'id': current.id, 'name': current.name, 'doc': current.doc})
            current = getattr(current, 'parent', None)
        suites.reverse()
        return suites


def parse_robot_output(xml_path):
    result = ExecutionResult(xml_path)
    collector = ScreenshotCollector()
    result.visit(collector)
    print(f"Found {len(collector.test_results)} test cases")
    return collector.test_results


class WordReportGenerator:
    def __init__(self, base_dir=None):
        self.base_dir = base_dir or PROJECT_ROOT
        self.temp_dir = None

    def format_action_message(self, message):
        if not message or "RobotTestExecutor: Executing action" not in message:
            return message
        try:
            action_match = re.search(r'action:$$([^$$]+)\]', message)
            page_match = re.search(r'page:$$([^$$]+)\]', message)
            module_match = re.search(r'module:$$([^$$]+)\]', message)
            element_match = re.search(r'element:$$([^$$]+)\]', message)
            params_match = re.search(r'with parameters:$$(.+?)$$(?:\s|$)', message)

            action = action_match.group(1) if action_match else 'Unknown Action'
            page = page_match.group(1) if page_match else 'Unknown Page'
            module = module_match.group(1) if module_match else 'Unknown Module'
            element = element_match.group(1) if element_match else None
            params = params_match.group(1) if params_match else None

            base_msg = f'Execute [{action}] in [{page}] page [{module}] module'
            if element:
                base_msg = f'{base_msg}, target element: [{element}]'
            if params:
                base_msg = f'{base_msg}, parameters: [{params}]'
            return base_msg
        except Exception as e:
            print(f"Error formatting message: {str(e)}")
            return message

    def process_base64_image(self, base64_data):
        if isinstance(base64_data, dict) and 'data' in base64_data:
            base64_data = base64_data['data']
        try:
            image_data = base64.b64decode(base64_data)
            return Image.open(io.BytesIO(image_data))
        except Exception:
            try:
                cleaned_data = base64_data.replace(" ", "").replace("\n", "")
                padding = 4 - (len(cleaned_data) % 4) if len(cleaned_data) % 4 != 0 else 0
                cleaned_data += "=" * padding
                image_data = base64.b64decode(cleaned_data)
                return Image.open(io.BytesIO(image_data))
            except Exception as e2:
                print(f"Failed to decode base64 data: {str(e2)}")
                return None

    def save_image_for_word(self, img, output_dir, index, test_name):
        try:
            os.makedirs(output_dir, exist_ok=True)
            safe_test_name = re.sub(r'[\\/*?:"<>|]', "_", test_name)
            filename = f"{safe_test_name}_{index}.jpg"
            filepath = os.path.join(output_dir, filename)

            if img.mode == 'RGBA':
                rgb_img = Image.new('RGB', img.size, (255, 255, 255))
                rgb_img.paste(img, mask=img.split()[3])
                rgb_img.save(filepath, 'JPEG', quality=95)
            else:
                if img.mode != 'RGB':
                    img = img.convert('RGB')
                img.save(filepath, 'JPEG', quality=95)
            return filepath
        except Exception as e:
            print(f"Error saving image: {str(e)}")
            return None

    def generate_report(self, xml_path=None, output_file=None):
        xml_path = xml_path or os.path.join(self.base_dir, 'reports', 'output.xml')
        output_file = output_file or os.path.join(self.base_dir, 'reports', 'test_evidences.docx')

        self.temp_dir = os.path.join(os.path.dirname(output_file), 'temp_screenshots')
        os.makedirs(self.temp_dir, exist_ok=True)

        test_results = parse_robot_output(xml_path)
        if not test_results:
            print('No test cases found')
            return None

        return self._generate_word_report(test_results, output_file)

    def _generate_word_report(self, test_results, output_file):
        doc = Document()

        # Cover Page
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("Robot Test Evidences")
        run.bold = True
        run.font.size = Pt(28)

        doc.add_paragraph("\nGenerated: " + datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
        doc.add_paragraph("\nThis report includes test results, details, and screenshots.")

        # Test Statistics
        total_cases = len(test_results)
        passed_cases = sum(1 for test in test_results if test["status"] == "PASS")
        failed_cases = sum(1 for test in test_results if test["status"] == "FAIL")
        skipped_cases = total_cases - (passed_cases + failed_cases)

        doc.add_heading("Test Summary", level=1)
        stats_table = doc.add_table(rows=2, cols=4)
        stats_table.style = 'Table Grid'

        header_cells = stats_table.rows[0].cells
        header_cells[0].text = 'Total Test Cases'
        header_cells[1].text = 'Passed'
        header_cells[2].text = 'Failed'
        header_cells[3].text = 'Skipped'

        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        value_cells = stats_table.rows[1].cells
        value_cells[0].text = str(total_cases)

        passed_paragraph = value_cells[1].paragraphs[0]
        passed_run = passed_paragraph.add_run(str(passed_cases))
        passed_run.font.color.rgb = RGBColor(0, 128, 0)  # Green
        passed_run.bold = True

        failed_paragraph = value_cells[2].paragraphs[0]
        failed_run = failed_paragraph.add_run(str(failed_cases))
        failed_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
        failed_run.bold = True

        skipped_paragraph = value_cells[3].paragraphs[0]
        skipped_run = skipped_paragraph.add_run(str(skipped_cases))
        skipped_run.font.color.rgb = RGBColor(255, 165, 0)  # Orange
        skipped_run.bold = True

        for row in stats_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph("\n")

        # Test Overview
        doc.add_heading("Test Case Overview", level=1)
        table = doc.add_table(rows=1, cols=5)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Test Case'
        hdr_cells[1].text = 'Start Time'
        hdr_cells[2].text = 'End Time'
        hdr_cells[3].text = 'Duration'
        hdr_cells[4].text = 'Result'

        for test in test_results:
            row_cells = table.add_row().cells
            row_cells[0].text = test['name']
            row_cells[1].text = test.get('start_time', '')
            row_cells[2].text = test.get('end_time', '')
            row_cells[3].text = f"{test.get('duration', '')} ms"
            result_text = 'Passed' if test['status'] == 'PASS' else 'Failed'
            row_cells[4].text = result_text

            run_status = row_cells[4].paragraphs[0].runs[0]
            if test['status'] == 'PASS':
                run_status.font.color.rgb = RGBColor(0, 128, 0)  # Green
            else:
                run_status.font.color.rgb = RGBColor(255, 0, 0)  # Red

        doc.add_paragraph("\n")

        # Detailed Test Results
        for i, test in enumerate(test_results):
            doc.add_page_break()

            doc.add_heading(f"Test Case: {test['name']}", level=1)

            p_status = doc.add_paragraph()
            p_status.add_run("Status: ").bold = True
            run_status = p_status.add_run("Passed" if test["status"] == "PASS" else "Failed")
            run_status.bold = True
            if test['status'] == 'PASS':
                run_status.font.color.rgb = RGBColor(0, 128, 0)  # Green
            else:
                run_status.font.color.rgb = RGBColor(255, 0, 0)  # Red

            doc.add_paragraph(f"Start Time: {test.get('start_time', '')}")
            doc.add_paragraph(f"End Time: {test.get('end_time', '')}")
            doc.add_paragraph(f"Duration: {test.get('duration', '')} ms")

            if test['doc']:
                doc.add_heading("Description", level=2)
                doc.add_paragraph(test['doc'])

            if test['steps']:
                doc.add_heading("Steps", level=2)
                for step_num, step in enumerate(test['steps'], 1):
                    p_step = doc.add_paragraph()
                    p_step.add_run(f"Step {step_num}: {step['name']}").bold = True
                    p_step.add_run(f"  [{step['status']}]")
                    if step['status'] != 'PASS':
                        p_step.runs[-1].font.color.rgb = RGBColor(255, 0, 0)

                    for msg in step['messages']:
                        formatted_msg = self.format_action_message(msg)
                        doc.add_paragraph(f"• {formatted_msg}", style='List Bullet')

            if test['base64_screenshots']:
                doc.add_heading("Screenshots", level=2)
                for screenshot_num, screenshot_info in enumerate(test['base64_screenshots'], 1):
                    step_info = screenshot_info.get('step', {})
                    description = step_info.get('description', "No description")
                    p_sc = doc.add_paragraph()
                    p_sc.add_run(f"Screenshot {screenshot_num}: {description}").bold = True

                    base64_data = screenshot_info['data']
                    img = self.process_base64_image(base64_data)
                    if img:
                        img_file = self.save_image_for_word(img, self.temp_dir, screenshot_num, test['name'])
                        if img_file and os.path.exists(img_file):
                            try:
                                doc.add_picture(img_file, width=Inches(6))
                            except Exception:
                                doc.add_paragraph("Failed to add screenshot.")
                    else:
                        doc.add_paragraph("Failed to process screenshot data.")
            else:
                doc.add_paragraph("No screenshots available.")

            doc.add_paragraph("\n" + "-" * 80 + "\n")

        try:
            doc.save(output_file)
            print(f'Report generated: {output_file}')
        except Exception as e:
            print(f'Error saving report: {str(e)}')
            return None

        self._cleanup_temp_dir()
        return output_file

    def _cleanup_temp_dir(self):
        try:
            if self.temp_dir and os.path.exists(self.temp_dir):
                shutil.rmtree(self.temp_dir)
        except Exception as e:
            print(f"Failed to clean temporary screenshot directory: {str(e)}")


def main():
    generator = WordReportGenerator()
    generator.generate_report()


if __name__ == "__main__":
    main()